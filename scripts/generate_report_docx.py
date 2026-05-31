# -*- coding: utf-8 -*-
"""生成三模型对照实验分析报告 (.docx)"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ---- 全局样式 ----
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')


def add_table(headers, rows):
    """添加带格式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
    return table


# ============================================================
# 正文
# ============================================================

# 标题
title = doc.add_heading('三模型对照实验分析报告', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')

# 一、评估指标说明
doc.add_heading('一、评估指标说明', level=2)

p = doc.add_paragraph()
p.add_run('本次实验采用 7 项精度指标和 1 项速度指标对模型进行全面评估。').bold = False

p = doc.add_paragraph()
p.add_run('mAP@0.5').bold = True
p.add_run('（mean Average Precision at IoU=0.5）是目标检测中最基础的精度指标，衡量模型在 IoU 阈值为 0.5 时的平均检测精度，数值越高说明模型越能正确找到目标。')
p.add_run('mAP@0.5:0.95').bold = True
p.add_run(' 则更为严格，它在 IoU 从 0.5 到 0.95 的多个阈值上取平均，反映了模型的定位精度——不仅要把目标找到，还要把框画准。')

p = doc.add_paragraph()
p.add_run('Precision').bold = True
p.add_run('（查准率）表示模型检测出的结果中有多少是真正正确的，Precision 低意味着误报多；')
p.add_run('Recall').bold = True
p.add_run('（查全率）表示所有真实目标中有多少被模型找到了，Recall 低意味着漏检多。两者往往此消彼长，因此引入 ')
p.add_run('F1 Score').bold = True
p.add_run(' 作为二者的调和均值，综合衡量检测质量。')

p = doc.add_paragraph()
p.add_run('每类 AP50').bold = True
p.add_run(' 将 mAP@0.5 拆分到 defect_2、defect_3、defect_6 三个瑕疵类别，可以定位模型在哪个具体类型上表现薄弱。')
p.add_run('FPPI').bold = True
p.add_run('（False Positive Per Image，每图误报数）在 398 张无瑕疵的负样本图片上统计，衡量模型的误报率，是工业部署场景的关键指标——产线上不能频繁误报。')

p = doc.add_paragraph()
p.add_run('推理速度').bold = True
p.add_run(' 以 OpenVINO 在 CPU 上测得的平均推理耗时（ms）和每秒处理帧数（FPS）衡量，直接决定模型能否满足产线实时检测需求。')

# 二、实验方法与模型
doc.add_heading('二、实验方法与模型', level=2)

p = doc.add_paragraph()
p.add_run('本次对照实验选取三个模型，在同一验证集（822 张图片，含 398 张负样本）上以相同参数（imgsz=640, conf=0.25, iou=0.45）进行评估，确保对比公平。')

add_table(
    ['模型', '说明'],
    [
        ['Official YOLOv8n', 'Ultralytics 官方 COCO 预训练权重，未做任何微调'],
        ['Baseline', '基于官方权重在鑫光板材数据集上微调 50 轮，mosaic=1.0'],
        ['CBAM+WIoU', '在 Baseline 基础上添加 CBAM 注意力模块（P3 层）并替换为 WIoU 损失函数，训练 100 轮'],
    ]
)

doc.add_paragraph('')

p = doc.add_paragraph()
p.add_run('Baseline 和 CBAM+WIoU 均基于相同的 YOLOv8n 架构和数据集配置（xinguang.yaml），唯一的区别是模型结构和损失函数的改动。')

# 三、实验结果
doc.add_heading('三、实验结果', level=2)

doc.add_heading('3.1 整体精度', level=3)

add_table(
    ['模型', 'mAP@0.5', 'mAP@0.5:0.95', 'Precision', 'Recall', 'F1'],
    [
        ['Official YOLOv8n', '0.0050', '0.0023', '0.0107', '0.0311', '0.0159'],
        ['Baseline', '0.5865', '0.3017', '0.6772', '0.5149', '0.5850'],
        ['CBAM+WIoU', '0.5712', '0.2818', '0.6660', '0.5138', '0.5801'],
    ]
)

doc.add_paragraph('')

p = doc.add_paragraph()
p.add_run('Official YOLOv8n 的 mAP50 仅 0.005，几乎无法检测任何瑕疵，说明 COCO 预训练权重对工业板材瑕疵场景没有迁移能力，微调是必要的。')

p = doc.add_paragraph()
p.add_run('对比 Baseline 和 CBAM+WIoU，后者在全部精度指标上均低于 Baseline：mAP50 低 2.6%，mAP50-95 低 6.6%，F1 低 0.8%。CBAM+WIoU 的改动不仅没有带来提升，反而造成了全面退化。值得注意的是 mAP50-95 的降幅（6.6%）远大于 mAP50 的降幅（2.6%），说明 CBAM+WIoU 的定位精度退化更为严重。')

doc.add_heading('3.2 每类检测能力', level=3)

p = doc.add_paragraph()
p.add_run('下表展示三个瑕疵类别的 AP50 和 Recall，用于定位模型在哪些具体类型上存在差异。')

add_table(
    ['模型', 'defect_2 AP50', 'defect_3 AP50', 'defect_6 AP50',
     'defect_2 Recall', 'defect_3 Recall', 'defect_6 Recall'],
    [
        ['Official YOLOv8n', '0.0128', '0.0000', '0.0021', '0.0933', '0.0000', '0.0000'],
        ['Baseline', '0.4167', '0.6894', '0.6533', '0.3785', '0.5349', '0.6312'],
        ['CBAM+WIoU', '0.3619', '0.6834', '0.6682', '0.3340', '0.6047', '0.6028'],
    ]
)

doc.add_paragraph('')

p = doc.add_paragraph()
p.add_run('三个模型都呈现出相同的规律：defect_2 最难检测，defect_3 和 defect_6 相对容易。Baseline 的 defect_2 AP50 仅 0.4167、Recall 仅 0.3785，意味着超过 60% 的 defect_2 被漏检，这是当前系统的核心瓶颈。')

p = doc.add_paragraph()
p.add_run('CBAM+WIoU 对不同类别产生了分化效果：defect_2 进一步恶化（AP50 下降 13.2%，Recall 下降 11.8%），defect_3 的 Recall 反而明显提升（从 0.5349 升至 0.6047，+13%），defect_6 基本持平。这说明 CBAM 注意力机制或 WIoU 损失函数对不同特征的瑕疵产生了截然不同的影响，两者捆绑在一起时，对 defect_3 的正面贡献被对 defect_2 的负面效应所抵消。')

doc.add_heading('3.3 负样本误报率', level=3)

add_table(
    ['模型', '负样本数', '有假框图片', '假框总数', 'FPPI', '状态'],
    [
        ['Official YOLOv8n', '398', '12', '14', '0.0352', 'PASS'],
        ['Baseline', '398', '26', '29', '0.0729', 'PASS'],
        ['CBAM+WIoU', '398', '29', '32', '0.0804', 'PASS'],
    ]
)

doc.add_paragraph('')

p = doc.add_paragraph()
p.add_run('三个模型的 FPPI 均低于 0.1 的工业部署阈值，误报率可控。但 CBAM+WIoU 的 FPPI（0.0804）比 Baseline（0.0729）高 10%，说明在无瑕疵图片上更容易产生误检，这在产线场景中意味着更高的停机复查成本。')

doc.add_heading('3.4 推理速度', level=3)

add_table(
    ['模型', '平均推理耗时 (ms)', '纯推理 FPS', '端到端 FPS'],
    [
        ['Official YOLOv8n', '9.89', '101.08', '51.84'],
        ['Baseline', '7.06', '141.57', '77.87'],
        ['CBAM+WIoU', '13.19', '75.81', '52.56'],
    ]
)

doc.add_paragraph('')

p = doc.add_paragraph()
p.add_run('CBAM 注意力模块带来了显著的计算开销：推理耗时从 Baseline 的 7.06ms 飙升至 13.19ms，增幅达 87%；FPS 从 141.57 降至 75.81，降幅 46%。这意味着 CBAM+WIoU 的推理速度不到 Baseline 的一半。在 DK2500 等边缘设备上，这一差距会更加明显。')

# 四、分析与结论
doc.add_heading('四、分析与结论', level=2)

p = doc.add_paragraph()
p.add_run('综合来看，CBAM+WIoU（Exp-02）的实验是失败的，但数据揭示了有价值的信息。')

p = doc.add_paragraph()
p.add_run('CBAM+WIoU 对 defect_2 的全面伤害是最大的问题。').bold = True
p.add_run(' defect_2 本身就是三类瑕疵中最难检测的（Baseline Recall 仅 0.3785），CBAM 的注意力机制可能在 defect_2 所在的特征尺度上产生了干扰，而 WIoU 的损失函数调整可能进一步削弱了对这类低对比度瑕疵的学习信号。defect_2 的平均框尺寸约 191×194px，属于中等尺寸目标，并非小目标问题，CBAM 针对小目标增强的设计初衷可能并不匹配。')

p = doc.add_paragraph()
p.add_run('WIoU 对 defect_3 可能有正面作用。').bold = True
p.add_run(' CBAM+WIoU 中 defect_3 的 Recall 提升了 13%，但由于两个变量同时改动，无法确定这是 WIoU 的贡献还是 CBAM 的贡献，甚至可能是两者的交互效应。这正是需要单变量实验来拆解的原因。')

p = doc.add_paragraph()
p.add_run('CBAM 的速度代价过高。').bold = True
p.add_run(' 87% 的推理耗时增长对于嵌入式部署来说是难以接受的，除非 CBAM 能带来显著的精度提升来证明其价值。目前的数据并不支持这一点。')

p = doc.add_paragraph()
p.add_run('基于以上分析，下一步建议按单变量原则拆解实验：先单独测试 WIoU 损失函数（Exp-03），验证其对 defect_3 的提升效果以及对 defect_2 的影响；再单独测试 CBAM 注意力模块（Exp-04），评估其精度收益是否值得速度代价。每轮先以 20 轮快速消融验证趋势，确认有效后再跑全量。')

# 保存
output_path = r"D:\files_1\PythonProject\XinGuang\docs\三模型对照实验分析报告.docx"
doc.save(output_path)
print(f"[DONE] 报告已保存至: {output_path}")
